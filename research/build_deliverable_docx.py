"""Generate the EU-project deliverable (.docx) on Layer-A KG construction + benchmarking.

Regenerable like the other build_* scripts. Run with the foodscholar env:

    /mnt/miniconda3/envs/foodscholar/bin/python scripts/build_deliverable_docx.py

Content is grounded in docs/methods_layer_a_bakeoff_brief.md and the implemented
bake-off harness; LLM-dependent results are labelled PRELIMINARY (no-GROQ run).
EU metadata fields are placeholders in [brackets] for the author to complete.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path("docs/deliverables/D_LayerA_KG_Construction_and_Benchmarking_v0.1.docx")

doc = Document()

# ---- base styling -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(text: str = "", *, italic: bool = False, bold: bool = False) -> None:
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = italic
    run.bold = bold


def bullets(items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def toc() -> None:
    """Insert a Word TOC field (updates on open / F9)."""
    par = doc.add_paragraph()
    run = par.add_run()
    fld_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    fld_text = run._r.makeelement(qn("w:t"), {})
    fld_text.text = "Right-click → Update Field to populate the table of contents."
    fld_end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(el)


# ============================================================ COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Layer A Knowledge Graph Construction and Benchmarking")
r.bold = True
r.font.size = Pt(22)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Entry-point index construction over the FoodOn ontology, and a "
                 "metric-driven method bake-off")
rs.italic = True
rs.font.size = Pt(13)
rs.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
doc.add_paragraph()

p("Project metadata (to be completed by the author):", bold=True)
table(
    ["Field", "Value"],
    [
        ["Project title", "[Project full title]"],
        ["Project acronym", "[ACRONYM]"],
        ["Grant Agreement No.", "[GA number]"],
        ["Deliverable No. & title", "[D x.y] — Layer A KG Construction and Benchmarking"],
        ["Work package / Task", "[WP x] / [Task x.y]"],
        ["Lead beneficiary", "[Partner]"],
        ["Dissemination level", "[PU / SEN / …]"],
        ["Deliverable type", "[R — Document, report]"],
        ["Due date / Submission date", "[MXX] / [YYYY-MM-DD]"],
        ["Version", "0.1 (draft)"],
        ["Authors", "[Author names, affiliations]"],
        ["Reviewers", "[Reviewer names]"],
    ],
)

p("Abstract", bold=True)
p("This deliverable documents the construction of Layer A of the project knowledge "
  "graph — a navigable entry-point (filter) index over a nutrition corpus, grounded "
  "in the FoodOn food ontology — and a reproducible benchmarking framework for "
  "selecting among competing construction methods. We characterise the central "
  "difficulty (FoodOn provides no single, browsable food hierarchy and trades "
  "faithfulness against navigability), describe six construction methods spanning "
  "purely structural to LLM-agentic approaches, and define an eight-metric scorecard "
  "that scores every method on identical inputs. We report a preliminary, "
  "LLM-free evaluation and the current implementation status, and we set out the "
  "remaining steps to a final, evidence-based method selection.", italic=True)

p("Keywords: knowledge graph; FoodOn; ontology projection; entry-point index; "
  "large language models; benchmarking; nutrition informatics.", italic=True)

doc.add_page_break()

# ============================================================ DOC CONTROL
h("Document control", level=1)
table(
    ["Version", "Date", "Author", "Description"],
    [
        ["0.1", "[YYYY-MM-DD]", "[Author]", "Initial draft: methods, benchmarking "
         "framework, preliminary LLM-free results."],
        ["", "", "", ""],
    ],
)
h("Table of contents", level=1)
toc()
doc.add_page_break()

# ============================================================ 1. EXEC SUMMARY
h("Executive summary", level=1)
p("Layer A is the entry-point layer of the knowledge graph: the set of recognisable "
  "categories a user selects to filter the corpus when they arrive with a specific "
  "question (for example, “does olive oil affect cardiovascular health?”). "
  "Its success criterion is coverage and nameability — for any food a user can name, "
  "a recognisable entry point should exist and be findable — rather than ontological "
  "completeness.")
p("The work is shaped by one structural fact: the FoodOn ontology does not provide a "
  "single, browsable food hierarchy. Foods are distributed across parallel "
  "classification axes; a substantial fraction of corpus-mentioned foods do not sit "
  "under the principal food product branch; and the is-a graph frequently fails to "
  "group foods as a human would. Consequently every construction method embodies a "
  "trade-off between faithfulness to FoodOn’s real structure (which is "
  "defensible and reproducible) and navigability (a shallow, recognisable tree).")
p("This deliverable (i) frames that trade-off, (ii) specifies six construction "
  "methods ranging from a purely structural backbone to an LLM agent that edits the "
  "ontology graph through a tool interface, (iii) defines a reproducible, "
  "eight-metric benchmarking harness that scores all methods on identical inputs and "
  "renders a single comparison scorecard, and (iv) reports the current "
  "implementation and a preliminary evaluation. The harness and a first set of "
  "methods are implemented and unit-tested; a complete, LLM-populated evaluation and "
  "the final method selection are the immediate next steps.")

# ============================================================ 2. INTRODUCTION
h("1  Introduction", level=1)
h("1.1  Purpose and scope", level=2)
p("This document covers the construction of Layer A (entry-point index) and the "
  "methodology for benchmarking competing construction methods. It does not cover "
  "upstream corpus ingestion and entity linking, nor the downstream thematic layer "
  "(Layer B), except where needed for context.")
h("1.2  Objective of Layer A", level=2)
p("Layer A is a faceted filter index rather than a complete taxonomy. The operative "
  "definition of success is: for any food a user would plausibly name, a recognisable "
  "entry point exists and is reachable in few interactions. Depth and ontological "
  "completeness are secondary to nameability and coverage.")
h("1.3  Relation to the wider knowledge graph", level=2)
p("Entry points (shelves) produced by Layer A anchor the corpus chunks attached to "
  "them and provide the substrate over which the thematic layer (Layer B) later "
  "discovers cross-cutting themes. The construction method chosen here therefore "
  "affects the entire downstream browsing experience.")

# ============================================================ 3. BACKGROUND
h("2  Background: FoodOn and the structural challenge", level=1)
p("FoodOn is an OWL ontology of foods and food-related entities. Two properties of "
  "FoodOn dominate the design space for Layer A:")
bullets([
    "No single food tree. Foods are organised along multiple parallel axes (for "
    "example by material vs. product, by process, and by consumer group). A "
    "single is-a projection therefore cannot reproduce an intuitive, mutually "
    "exclusive set of food groups.",
    "Partial and counter-intuitive grouping. A significant fraction of "
    "corpus-mentioned foods do not lie under the principal food product branch, and "
    "the is-a graph often does not connect a specific food to the human category a "
    "user expects (for example, a specific fruit may not connect to a usable "
    "“fruit” food node).",
])
p("These properties create a fundamental tension between two desirable properties of "
  "the constructed index:")
table(
    ["Property", "Meaning", "Cost"],
    [
        ["Faithfulness", "The hierarchy is FoodOn’s real (is-a or relational) "
         "structure.", "May be deep, wide, or sparsely populated for browsing."],
        ["Navigability", "A shallow, recognisable, well-populated tree.", "May "
         "require departing from FoodOn’s real structure."],
    ],
)
p("The benchmarking framework in Section 4 exists precisely to measure this "
  "trade-off rather than to argue it qualitatively.")

# ============================================================ 4. METHODS
h("3  Construction methods", level=1)
h("3.1  Design space", level=2)
p("Each method is a choice on two independent axes: (a) the source of hierarchy and "
  "membership — FoodOn’s structure versus LLM semantic judgement; and (b) the "
  "source of coverage — top-down (start at roots, prune) versus bottom-up (start at "
  "corpus-mentioned leaves). Bottom-up construction guarantees coverage by "
  "construction; structural membership preserves faithfulness.")
h("3.2  Method catalogue", level=2)
table(
    ["ID", "Method", "Hierarchy source", "Faithfulness", "Notes"],
    [
        ["0", "Top-down prune (baseline)", "is-a, top-down", "High", "Reference; "
         "tends to a flat, sparsely-tiered result."],
        ["1a", "Structural auto-backbone", "is-a + support", "High", "Fixed backbone; "
         "corpus support decorates it."],
        ["1a+", "Backbone + controlled expansion", "is-a + support", "High", "Opens "
         "deeper tiers only where supported; fan-out and depth capped; single-child "
         "chains skipped. Strong faithful-and-navigable baseline."],
        ["1b", "LLM-proposed backbone", "LLM names → is-a fill", "Medium", "LLM "
         "proposes top-level categories, each resolved to a real FoodOn id."],
        ["G", "Bottom-up LLM grouping", "LLM by label", "Low", "Groups foods by label, "
         "ignoring is-a; navigable but not reproducible from structure."],
        ["A", "Agentic MCP editor", "is-a (+ relations)", "High", "LLM agent traverses "
         "the FoodOn support graph and makes local keep/collapse/reparent decisions "
         "through a tool interface; never fabricates parents."],
    ],
)
h("3.3  The controlled-expansion baseline (1a+)", level=2)
p("Method 1a+ keeps a structural backbone and then recursively opens deeper FoodOn "
  "tiers only where corpus support exceeds a floor, capping fan-out per parent and "
  "overall depth and skipping low-value single-child chains. Every displayed node "
  "remains a real FoodOn identifier, so the method is simultaneously faithful, "
  "bottom-up (covered), and navigable. It is the principal baseline that the agentic "
  "method must outperform.")
h("3.4  The agentic MCP method (A)", level=2)
p("The agentic method tests whether an LLM that reasons over the real ontology graph "
  "can outperform the mechanical rules of 1a+. An agent traverses the FoodOn support "
  "graph top-down; at each node it is shown a local “lens” (the node, its "
  "parent, supported children with their support counts, and candidate non-is-a "
  "relation bridges) and returns one local action over real edges:")
bullets([
    "KEEP — the node is a recognisable category and becomes a shelf;",
    "COLLAPSE — the node is redundant with its parent; its children are lifted up;",
    "REPARENT — the node is an organisational artefact; its children are lifted to "
    "the parent.",
])
p("Because the agent acts only on real FoodOn edges, membership remains is-a-faithful "
  "and every decision is accompanied by a recorded rationale (an audit trail). A "
  "throwaway relation index extracted from the FoodOn OWL exposes non-is-a relations "
  "(for example derives from, member of, has ingredient) to the agent; acting on "
  "these relations to bridge gaps in the is-a graph is specified but deferred to a "
  "subsequent increment (Section 7). The agent operates through a manual "
  "tool-invocation protocol because the underlying LLM client does not provide native "
  "function calling.")

# ============================================================ 5. BENCHMARKING
h("4  Benchmarking methodology", level=1)
h("4.1  Principle", level=2)
p("Methods are not judged qualitatively. Each method emits a common representation "
  "(a directed tree with per-leaf placement records), and a fixed set of metrics is "
  "computed identically over every method on the same corpus and ontology snapshot. "
  "The output is a single scorecard plus, for the LLM methods, a per-decision audit "
  "trail.")
h("4.2  Metrics", level=2)
table(
    ["Metric", "What it captures", "Definition (summary)"],
    [
        ["Coverage", "Are named foods reachable?", "Fraction of mentioned leaves "
         "homed under some node. Near 1.0 for any bottom-up method, hence not the "
         "discriminator."],
        ["Specificity", "Are foods placed precisely?", "Mean/median is-a distance "
         "from each leaf to its home node. The discriminator coverage is not: "
         "separates a flat blob from a well-tiered tree."],
        ["Findability", "Effort to reach a food", "Held-out food queries → clicks "
         "from root to home; median, 90th percentile, and fraction within K."],
        ["Nameability", "Are labels recognisable?", "Fraction of a sampled set of "
         "shelf labels judged recognisable by an LLM (or human)."],
        ["Fan-out", "Navigation width", "Maximum / median children per node."],
        ["Depth", "Navigation depth", "Maximum / median depth of placed foods."],
        ["Faithfulness", "Is the structure real?", "Fraction of leaf placements that "
         "are is-a / non-is-a-relation / fabricated."],
        ["Reproducibility & cost", "Stability and price", "Number of LLM calls; "
         "run-to-run set overlap (Jaccard)."],
    ],
)
h("4.3  Procedure", level=2)
numbered([
    "Build a shared support layer (mentioned food leaves → chunk evidence) once "
    "from a single corpus and FoodOn snapshot.",
    "Run each method to produce its common representation.",
    "Sample a held-out, stratified query set of foods (common and rare) for the "
    "findability metric.",
    "Compute the eight metrics per method and assemble the scorecard.",
    "Render each method’s tree for inspection and, for LLM methods, export the "
    "decision audit.",
])

# ============================================================ 6. IMPLEMENTATION
h("5  Implementation", level=1)
p("The harness is implemented as a focused Python package with full unit-test "
  "coverage and is exercised from a single evaluation notebook that renders the "
  "scorecard above side-by-side method trees. Construction methods already present in "
  "the project (baseline, structural backbone, controlled expansion, LLM backbone, "
  "and bottom-up grouping) are adapted into the common representation; the agentic "
  "method is implemented as a new sub-package.")
table(
    ["Component", "Responsibility"],
    [
        ["Common representation", "A single result type (tree edges, labels, counts, "
         "per-leaf home and placement type/distance) emitted by every method."],
        ["Metric functions", "Pure functions computing the eight metrics over the "
         "common representation."],
        ["Scorecard", "Assembles one row per method and renders a comparison table."],
        ["Adapters", "Wrap each existing method’s output into the common "
         "representation."],
        ["Agentic sub-package", "Relation index (from OWL), per-node support roll-up, "
         "read-only graph tools, and the DFS decision loop."],
        ["Evaluation notebook", "Single entry point: runs all methods, renders the "
         "scorecard, trees, and audits."],
    ],
)
p("The construction and evaluation code is unit-tested (26 tests for the harness and "
  "agentic sub-package at the time of writing) and is regenerable from source build "
  "scripts. Earlier exploratory notebooks have been archived behind the single "
  "evaluation notebook to remove duplication.")

# ============================================================ 7. RESULTS
h("6  Preliminary results", level=1)
p("The figures below are PRELIMINARY and were produced in an LLM-free configuration "
  "(no API key present). In that configuration the LLM-dependent columns "
  "(nameability; the agentic method; the live grouping method, which degrades to one "
  "shelf per leaf) are not populated, and the faithfulness discriminator cannot "
  "differentiate methods. They are reported to demonstrate the harness end-to-end and "
  "to surface methodological findings, not to select a method.", italic=True)
h("6.1  Preliminary scorecard (LLM-free run)", level=2)
table(
    ["Method", "Coverage", "Find. median (clicks)", "Find. ≤ 3", "Max fan-out", "Max depth"],
    [
        ["0 — Baseline", "1.00", "2", "1.00", "111", "3"],
        ["2 — Structural cut", "1.00", "3", "0.92", "85", "4"],
        ["1a — Auto backbone", "1.00", "2", "0.99", "30", "2"],
        ["1a+ — Controlled expansion", "1.00", "5", "0.16", "12", "6"],
        ["3 — Multi-facet", "1.00", "1", "0.99", "10", "1"],
        ["G — Grouping (LLM-free)", "1.00", "1", "1.00", "2951", "1"],
    ],
)
p("Corpus snapshot for this run: of the order of 13,000 corpus chunks, of which "
  "approximately 4,200 mention at least one food; approximately 2,950 distinct "
  "food leaves; FoodOn snapshot of the order of 39,000 terms.")
h("6.2  Agentic method on real data", level=2)
p("Executed against the full ontology with a deterministic placeholder for the LLM "
  "decision (constant KEEP), the agentic loop completed in well under one second, "
  "respected the depth cap, produced approximately 200 internal nodes, and homed "
  "1,841 of 2,951 mentioned leaves. The remaining approximately 38% are foods that "
  "do not lie under the principal food product branch — precisely the gap that the "
  "deferred relation-bridging increment is designed to close.")
h("6.3  Methodological findings", level=2)
bullets([
    "Coverage does not discriminate. Because any ancestor in the tree counts as a "
    "home, coverage is approximately 1.0 for every method. The specificity metric was "
    "introduced for this reason and is the discriminator between a flat blob and a "
    "well-tiered tree.",
    "Findability penalises depth by construction. Because a leaf is homed to its "
    "deepest kept ancestor, intrinsically deep trees (such as 1a+) score lower on the "
    "fraction-within-three-clicks measure than a user who stops at a recognisable "
    "higher tier would experience. This is an interpretation caveat to be addressed "
    "in the metric refinement.",
    "LLM-dependent metrics require a live model. Nameability, the live grouping "
    "method, and the agentic method in the notebook are gated on model access; a "
    "complete evaluation requires a configured LLM.",
])

# ============================================================ 8. DISCUSSION
h("7  Discussion", level=1)
p("The preliminary run confirms the central tension quantitatively in miniature: the "
  "shallow methods (multi-facet, 1a) reach foods in few clicks but at the cost of "
  "wide, generic top tiers, while the controlled-expansion method (1a+) places foods "
  "more specifically at the cost of depth. The genuinely informative comparison — "
  "specificity and nameability against depth and fan-out, with faithfulness held to "
  "real FoodOn structure — requires the LLM-populated run and is the immediate next "
  "step. The agentic method is positioned as the test of whether context-aware, "
  "per-node LLM judgement over the real graph improves on the mechanical rules of "
  "1a+ without sacrificing faithfulness.")

# ============================================================ 9. LIMITATIONS
h("8  Limitations", level=1)
bullets([
    "Results reported here are LLM-free and therefore partial.",
    "The relation-bridging capability of the agentic method (acting on non-is-a "
    "FoodOn relations to close the not-under-food-product gap) is specified but not "
    "yet implemented.",
    "The non-is-a relation index is a prototype loaded outside the production "
    "ontology pipeline; promoting it to the production loader is future work.",
    "Two metrics require refinement (a discriminating coverage variant and a "
    "tier-aware findability variant), as identified in Section 6.3.",
    "LLM-based methods are not bit-for-bit reproducible; the reproducibility metric "
    "quantifies, but does not remove, this variation.",
])

# ============================================================ 10. CONCLUSIONS
h("9  Conclusions and next steps", level=1)
p("A reproducible benchmarking framework for Layer A construction has been "
  "established, together with a catalogue of six construction methods spanning purely "
  "structural to LLM-agentic approaches, and the agentic method’s faithful core "
  "has been implemented and tested. Method selection will be evidence-based, driven "
  "by the scorecard rather than by argument. The immediate next steps are:")
numbered([
    "Execute the full evaluation with a configured LLM to populate the nameability, "
    "faithfulness, grouping, and agentic columns of the scorecard.",
    "Implement the agentic relation-bridging increment to close the "
    "not-under-food-product coverage gap, with each bridged edge logged by relation "
    "type.",
    "Refine the coverage and findability metrics per Section 6.3.",
    "Select the balanced construction method from the completed scorecard.",
    "Promote the selected method to the production construction path and, if "
    "warranted, the relation index to the production ontology loader.",
])

# ============================================================ APPENDICES
doc.add_page_break()
h("Appendix A  Metric definitions (summary)", level=1)
bullets([
    "Coverage = |homed mentioned leaves| / |mentioned leaves|.",
    "Specificity = mean (and median) over homed leaves of the number of is-a steps "
    "from the leaf to its home node (0 if the leaf is itself a node).",
    "Findability = over a held-out query set, the distribution (median, 90th "
    "percentile, fraction ≤ K) of root-to-home depths.",
    "Nameability = fraction of a sampled label set judged recognisable.",
    "Fan-out / Depth = maximum and median children per node / depth of homed leaves.",
    "Faithfulness = fractional split of leaf placements into is-a, non-is-a-relation, "
    "and fabricated.",
    "Reproducibility & cost = LLM-call count; Jaccard overlap of node sets across two "
    "runs.",
])
h("Appendix B  Key configuration parameters (defaults)", level=1)
table(
    ["Parameter", "Default", "Role"],
    [
        ["Minimum support to open a tier", "25 chunks", "Controls visibility of "
         "deeper tiers (1a+, agentic)."],
        ["Maximum fan-out per parent", "12", "Bounds navigation width."],
        ["Maximum depth", "6", "Bounds navigation depth."],
        ["Findability query set size", "~100", "Stratified common/rare foods."],
        ["Nameability sample size", "25", "Labels judged per method."],
    ],
)
h("Appendix C  Software manifest", level=1)
bullets([
    "Harness package: result representation, metric functions, scorecard renderer, "
    "and method adapters.",
    "Agentic sub-package: OWL relation index, support roll-up, read-only graph tools, "
    "and the DFS decision loop.",
    "Single evaluation notebook (regenerable from a build script) producing the "
    "scorecard, method trees, and decision audits.",
    "Unit tests: 26 covering the harness and agentic components.",
    "Design references: the Layer-A method bake-off brief and the two implementation "
    "plans (harness; agentic method).",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT} ({OUT.stat().st_size/1024:.0f} KB, {len(doc.paragraphs)} paragraphs)")
